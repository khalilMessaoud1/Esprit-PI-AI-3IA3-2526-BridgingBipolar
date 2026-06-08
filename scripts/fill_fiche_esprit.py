"""Fill ESPRIT GitHub submission Word form."""
from pathlib import Path
import shutil
from docx import Document

SRC = Path(r"c:\Users\LOQ\Downloads\Fiche_Soumission_GitHub_ESPRIT.docx")
DST = Path(__file__).resolve().parents[1] / "docs" / "Fiche_Soumission_GitHub_ESPRIT.docx"
DOWNLOADS_OUT = Path(r"c:\Users\LOQ\Downloads\Fiche_Soumission_GitHub_ESPRIT_BridgingBipolar.docx")

REPLACEMENTS = {
    "Classe": "3IA3",
    "Groupe": "Alpha",
    "Tuteur": "Dr. Jihene HLEL, Mrs. Wided ASKRI, Mr. Fedi BACCAR",
    "Adresse e-mail": "khalil.messaoud@esprit.tn",
    "Nom du projet": "Esprit-PI-AI-3IA3-2526-BridgingBipolar",
    "Lien GitHub": "https://github.com/khalilMessaoud1/Esprit-PI-AI-3IA3-2526-BridgingBipolar",
    "Technologies utilis": (
        "Next.js 14, NestJS, PostgreSQL, Prisma, Python (FastAPI), "
        "Ollama, Docker, Tailwind, GraphRAG"
    ),
    "Type de projet": "Web + IA",
    "Commande de lancement": (
        r"npm install → .\scripts\setup-env.ps1 → npm run dev:deps → .\scripts\start-dev.ps1"
    ),
    "Temps d": "~10 min (web + API) · 20–40 min (stack IA complète)",
    "Lien de d": "Non disponible",
    "Date de soumission": "08/06/2026",
    "Signature de l": "Khalil MESSAOUD",
}

SKIP_LABELS = {"champ", "élément à vérifier", "technologie / outil", "réponse attendue"}


def fill_doc(path: Path) -> None:
    doc = Document(path)

    # Table 0 — équipe (ligne 0 = en-têtes, ligne 1 = valeurs)
    if doc.tables:
        team = doc.tables[0]
        if len(team.rows) >= 2 and len(team.rows[1].cells) >= 4:
            team.rows[1].cells[0].text = "3IA3"
            team.rows[1].cells[1].text = "Alpha"
            team.rows[1].cells[2].text = "Dr. Jihene HLEL, Mrs. Wided ASKRI, Mr. Fedi BACCAR"
            team.rows[1].cells[3].text = "khalil.messaoud@esprit.tn"

    for table_idx, table in enumerate(doc.tables):
        if table_idx == 0:
            continue
        for row in table.rows:
            cells = row.cells
            if len(cells) < 2:
                continue
            label = cells[0].text.strip()
            label_l = label.lower()
            if label_l in SKIP_LABELS:
                continue
            for key, val in REPLACEMENTS.items():
                if key.lower() in label_l and key not in ("Date de soumission", "Signature de l"):
                    cells[1].text = val
                    break
            if len(cells) >= 3 and table_idx == 2:
                yes_items = [
                    "nom du d",
                    "dépôt public",
                    "readme",
                    "gitignore",
                    "env.example",
                    "docs/",
                    "sans aide",
                    "base de donn",
                    "credential",
                ]
                if any(x in label_l for x in yes_items):
                    cells[-1].text = "Confirmé" if "credential" in label_l else "Oui"
            if table_idx == 3:
                if "node" in label_l:
                    cells[1].text = "20+"
                elif "python" in label_l:
                    cells[1].text = "3.11+"
                elif "docker" in label_l:
                    cells[1].text = "Requis"
                elif label_l == "autre":
                    cells[1].text = "Ollama (optionnel — stack IA complète)"

    # Table 4 — date + signature (ligne 0 = en-têtes, ligne 1 = valeurs)
    if len(doc.tables) >= 5:
        sig = doc.tables[4]
        if len(sig.rows) >= 2:
            sig.rows[1].cells[0].text = "08/06/2026"
            if len(sig.rows[1].cells) >= 2:
                sig.rows[1].cells[1].text = "Khalil MESSAOUD"

    doc.save(path)


def main() -> None:
    DST.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(SRC, DST)
    fill_doc(DST)
    shutil.copy2(DST, DOWNLOADS_OUT)
    print(f"Saved: {DST}")
    print(f"Saved: {DOWNLOADS_OUT}")


if __name__ == "__main__":
    main()
